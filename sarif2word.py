import json
import sys

from docx import Document
from docx.shared import Pt


def generate_word_from_json(json_data, output_file):
    # 创建一个Word文档对象
    doc = Document()

    # 添加主标题
    title = json_data.get("title", "")
    if title:
        doc.add_heading(title, level=0)  # 0级标题是最大的标题

    # 添加摘要
    summary = json_data.get("summary", "")
    if summary:
        doc.add_paragraph(summary)
        doc.add_paragraph("")  # 添加空行分隔

    # 处理数据列表
    data_list = json_data.get("data", [])
    for item in data_list:
        # 添加项目标题
        heading = item.get("heading", "")
        if heading:
            doc.add_heading(heading, level=1)

        # 添加描述作为编号列表的第1项
        description = item.get("description", "")
        if description:
            p = doc.add_paragraph(style='List Number')
            p.add_run(description).font.size = Pt(12)

        # 添加位置作为编号列表的第2项
        location = item.get("location", "")
        if location:
            p = doc.add_paragraph(style='List Number')
            p.add_run(location).font.size = Pt(12)

        # 在每个项目后添加一个空段落作为分隔
        doc.add_paragraph("")

    # 保存Word文档
    doc.save(output_file)
    print(f"Word文档已成功生成：{output_file}")


def generate(path, commit_sha, output_file):
    import json
    try:
        with open(path, 'r', encoding='utf-8') as file:
            data = json.load(file)
            data_detail = []
            trim_data = data['runs'][0]['results']
            for x in trim_data:
                item = {
                    "heading": x['ruleId'],
                    "location": "location: " +
                                x['locations'][0]['physicalLocation']['artifactLocation']['uri'] + ":" +
                                str(x['locations'][0]['physicalLocation']['region']['startLine']),
                    "description": x['message']['text']
                }
                data_detail.append(item)

            final_data = {
                "title": "Code-scan report",
                "summary": "The commit sha is: " + commit_sha,
                "data": data_detail
            }
            generate_word_from_json(final_data, output_file)
    except Exception as e:
        print(e)

if __name__ == "__main__":
    generate("./results/java.sarif", sys.argv[1],
             "./results/output.docx")
